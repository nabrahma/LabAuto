from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import json
import os
import io
import re
import base64

# Third-party imports
import google.generativeai as genai
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfplumber

app = Flask(__name__)
CORS(app)


def configure_matlab_style():
    """Configure Matplotlib to mimic MATLAB default aesthetics."""
    plt.style.use('default')
    plt.rcParams['font.family'] = 'sans-serif'
    plt.rcParams['font.sans-serif'] = ['Helvetica', 'Arial', 'DejaVu Sans']
    plt.rcParams['font.size'] = 10
    plt.rcParams['axes.linewidth'] = 1.0
    plt.rcParams['xtick.direction'] = 'in'
    plt.rcParams['ytick.direction'] = 'in'
    plt.rcParams['xtick.top'] = True
    plt.rcParams['ytick.right'] = True
    plt.rcParams['grid.alpha'] = 0.5
    plt.rcParams['grid.linestyle'] = ':'
    plt.rcParams['figure.facecolor'] = 'white'
    plt.rcParams['axes.facecolor'] = 'white'
    plt.rcParams['axes.prop_cycle'] = plt.cycler(
        color=['#0072BD', '#D95319', '#EDB120', '#7E2F8E', '#77AC30', '#4DBEEE', '#A2142F']
    )


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return '\n'.join(text_parts)


def extract_text_from_docx(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    return '\n'.join(text_parts)


def split_questions(text: str) -> list:
    """Split input text into individual questions based on numbering patterns."""
    lines = text.split('\n')
    questions = []
    current_q = []
    
    # Pattern to detect question start: 1), 1., Q1, etc.
    q_pattern = re.compile(r'^\s*(\d+)\s*[\.\)]\s*')
    
    for line in lines:
        match = q_pattern.match(line)
        if match and len(current_q) > 0:
            # Save previous question
            q_text = '\n'.join(current_q).strip()
            if len(q_text) > 20:
                questions.append(q_text)
            current_q = [line]
        else:
            current_q.append(line)
    
    # Don't forget the last question
    if current_q:
        q_text = '\n'.join(current_q).strip()
        if len(q_text) > 20:
            questions.append(q_text)
    
    # If no questions found, return the whole text
    if not questions:
        questions = [text.strip()]
    
    return questions


def call_gemini_single(question_text: str, question_num: int) -> dict:
    """Call Gemini for a single question with robust error handling."""
    api_key = os.environ.get('GEMINI_API_KEY')
    if not api_key:
        raise ValueError("GEMINI_API_KEY environment variable not set")
    
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-2.0-flash')
    
    system_prompt = """You are a MATLAB expert. For the given lab question, provide:

1. MATLAB code that solves the problem (complete and executable, with hardcoded values)
2. Equivalent Python code using matplotlib/numpy to generate the same plot
3. A brief conclusion (2-3 sentences)

CRITICAL FORMATTING RULES:
- Output ONLY a JSON object with keys: matlab_code, python_plotting_code, conclusion
- Use \\n for newlines in code strings (not actual line breaks within the JSON string values)
- Escape all backslashes and quotes properly in JSON
- The python_plotting_code must end with: plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
- Do NOT use markdown code fences
- Keep code concise"""

    prompt = f"Question {question_num}: {question_text[:1500]}"

    try:
        response = model.generate_content(
            [system_prompt, prompt],
            generation_config=genai.types.GenerationConfig(
                temperature=0.2,
                max_output_tokens=4096,
            )
        )
        
        response_text = response.text.strip()
        
        # Clean up response
        response_text = re.sub(r'^```(?:json)?\s*', '', response_text)
        response_text = re.sub(r'\s*```$', '', response_text)
        
        # Try to find JSON object
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            response_text = json_match.group()
        
        # Parse JSON with error recovery
        try:
            result = json.loads(response_text)
        except json.JSONDecodeError:
            # Try to fix common JSON issues
            fixed = response_text
            fixed = re.sub(r',\s*}', '}', fixed)  # Remove trailing commas
            fixed = re.sub(r',\s*]', ']', fixed)
            try:
                result = json.loads(fixed)
            except:
                # Last resort: extract parts manually
                matlab_match = re.search(r'"matlab_code"\s*:\s*"([^"]*(?:\\"[^"]*)*)"', response_text)
                python_match = re.search(r'"python_plotting_code"\s*:\s*"([^"]*(?:\\"[^"]*)*)"', response_text)
                conclusion_match = re.search(r'"conclusion"\s*:\s*"([^"]*(?:\\"[^"]*)*)"', response_text)
                
                result = {
                    'matlab_code': matlab_match.group(1) if matlab_match else '% Code parsing error',
                    'python_plotting_code': python_match.group(1) if python_match else '',
                    'conclusion': conclusion_match.group(1) if conclusion_match else 'Conclusion parsing error'
                }
        
        # Unescape the code
        matlab_code = result.get('matlab_code', '% No code generated')
        python_code = result.get('python_plotting_code', '')
        conclusion = result.get('conclusion', 'No conclusion.')
        
        # Fix escaped newlines
        matlab_code = matlab_code.replace('\\n', '\n')
        python_code = python_code.replace('\\n', '\n')
        
        return {
            'matlab_code': matlab_code,
            'python_plotting_code': python_code,
            'conclusion': conclusion
        }
        
    except Exception as e:
        return {
            'matlab_code': f'% Error: {str(e)[:100]}',
            'python_plotting_code': '',
            'conclusion': f'Error processing question: {str(e)[:50]}'
        }


def generate_graph(python_code: str) -> bytes:
    """Execute Python plotting code and return PNG bytes."""
    configure_matlab_style()
    buffer = io.BytesIO()
    
    if not python_code or len(python_code.strip()) < 20:
        # Generate a placeholder graph
        fig, ax = plt.subplots(figsize=(8, 6))
        ax.text(0.5, 0.5, 'Graph generation pending', 
                ha='center', va='center', fontsize=14, color='gray')
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        ax.axis('off')
        plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
        buffer.seek(0)
        plt.close('all')
        return buffer.read()
    
    exec_globals = {
        'np': np,
        'plt': plt,
        'buffer': buffer,
        '__builtins__': {
            'range': range, 'len': len, 'abs': abs, 'min': min, 'max': max,
            'sum': sum, 'round': round, 'int': int, 'float': float,
            'str': str, 'list': list, 'tuple': tuple, 'dict': dict, 'print': print,
        }
    }
    
    try:
        plt.close('all')
        exec(python_code, exec_globals)
        
        if buffer.tell() == 0:
            plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight',
                       facecolor='white', edgecolor='none')
        
        buffer.seek(0)
        return buffer.read()
    except Exception as e:
        plt.close('all')
        fig, ax = plt.subplots(figsize=(8, 6))
        error_msg = str(e)[:120]
        ax.text(0.5, 0.5, f'Graph Error:\n{error_msg}',
               ha='center', va='center', fontsize=9, color='red',
               wrap=True)
        ax.axis('off')
        
        error_buffer = io.BytesIO()
        plt.savefig(error_buffer, format='png', dpi=150, bbox_inches='tight')
        error_buffer.seek(0)
        plt.close('all')
        return error_buffer.read()


def assemble_document(questions_data: list, student_name: str, roll_number: str, lab_number: str) -> bytes:
    """
    Assemble a document matching the user's college template format.
    Header with name/roll/lab, red question text, code block, graph, output label.
    """
    doc = Document()
    
    # Header section
    if student_name:
        p = doc.add_paragraph()
        run = p.add_run(f"Name – {student_name}")
        run.bold = True
        run.font.color.rgb = RGBColor(139, 0, 0)  # Dark red
        run.font.size = Pt(12)
    
    if roll_number:
        p = doc.add_paragraph()
        run = p.add_run(roll_number)
        run.font.color.rgb = RGBColor(139, 0, 0)
        run.font.size = Pt(12)
    
    if lab_number:
        p = doc.add_paragraph()
        run = p.add_run(f"WCT LAB {lab_number}")
        run.font.color.rgb = RGBColor(139, 0, 0)
        run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    for item in questions_data:
        # Question text in RED
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"{item['question_num']})  {item['question'][:400]}")
        q_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        q_run.font.size = Pt(11)
        
        doc.add_paragraph()  # Spacing
        
        # MATLAB Code in black, code-style font
        code_lines = item['matlab_code'].split('\n')
        for line in code_lines:
            p = doc.add_paragraph()
            # Check if line has special formatting (underline for function calls)
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
        
        doc.add_paragraph()  # Spacing
        
        # Output label
        output_para = doc.add_paragraph()
        output_run = output_para.add_run("Output:")
        output_run.font.color.rgb = RGBColor(255, 0, 0)
        output_run.font.size = Pt(11)
        
        # Graph
        if item['graph_bytes']:
            doc.add_picture(io.BytesIO(item['graph_bytes']), width=Inches(5))
        
        doc.add_paragraph()  # Spacing
        
        # Page break between questions (except last)
        if item != questions_data[-1]:
            doc.add_page_break()
    
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    return output_buffer.read()


@app.route('/', methods=['GET'])
def health():
    return jsonify({
        'status': 'ok', 
        'service': 'LabAuto API', 
        'version': '2.1',
        'features': ['multi-question', 'template-format', 'matlab-graphs']
    })


@app.route('/generate', methods=['POST', 'OPTIONS'])
def generate():
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.get_json()
        
        question_text = data.get('question_text', '')
        file_data = data.get('file_data')
        file_type = data.get('file_type', '')
        student_name = data.get('student_name', '')
        roll_number = data.get('roll_number', '')
        lab_number = data.get('lab_number', '')
        
        # Extract text from file if provided
        if file_data and not question_text:
            file_bytes = base64.b64decode(file_data)
            if file_type == 'pdf':
                question_text = extract_text_from_pdf(file_bytes)
            elif file_type == 'docx':
                question_text = extract_text_from_docx(file_bytes)
        
        if not question_text:
            return jsonify({'error': 'No question text provided'}), 400
        
        # Split into individual questions
        questions = split_questions(question_text)
        
        # Limit to 4 questions max to avoid timeout
        if len(questions) > 4:
            questions = questions[:4]
        
        # Process each question
        questions_data = []
        for i, q in enumerate(questions, 1):
            # Call Gemini
            ai_response = call_gemini_single(q, i)
            
            # Generate graph
            graph_bytes = generate_graph(ai_response['python_plotting_code'])
            
            questions_data.append({
                'question_num': i,
                'question': q,
                'matlab_code': ai_response['matlab_code'],
                'graph_bytes': graph_bytes,
                'conclusion': ai_response['conclusion']
            })
        
        # Assemble document
        doc_bytes = assemble_document(questions_data, student_name, roll_number, lab_number)
        
        return send_file(
            io.BytesIO(doc_bytes),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='lab_report.docx'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
